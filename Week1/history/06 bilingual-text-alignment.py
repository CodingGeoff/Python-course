
# %% [Cell 1]

# Cell 1
# uv pip install pandas openpyxl python-docx



# %% [Cell 2]

# Cell 2
import os
import glob

def find_and_preview_file():
    """
    Scans the current directory for .txt files containing 'bilingual'.
    Returns the filename if found, otherwise creates a demo file.
    """
    # 1. Define the search pattern
    # The asterisk (*) is a wildcard. "*bilingual*.txt" means:
    # "Any characters" + "bilingual" + "Any characters" + ".txt"
    search_pattern = "*bilingual*.txt"
    
    # 2. Use glob to find all matching files in the folder
    found_files = glob.glob(search_pattern)
    
    target_filename = ""

    # 3. Decision Logic
    if len(found_files) > 0:
        # If we found files, we pick the first one automatically
        target_filename = found_files[0]
        print(f"✅ Auto-Detection Successful!")
        print(f"📂 Found {len(found_files)} file(s). We will process: '{target_filename}'")
        
        # If there are multiple files, list them for user awareness
        if len(found_files) > 1:
            print(f"   (Other files ignored: {found_files[1:]})")
            
    else:
        # If NO file is found, we create a dummy one so the code doesn't crash
        print("❌ No text file containing 'bilingual' was found.")
        print("⚠️ Creating a 'demo_bilingual_test.txt' for demonstration purposes...")
        
        target_filename = "demo_bilingual_test.txt"
        dummy_content = """
        Contract Agreement for Service
        服务合同协议书
        
        This Agreement is made on January 23, 2026.
        本协议于 2026 年 1 月 23 日 签订。
        
        Party A: Tech Solutions Inc. (The "Client")
        甲方：Tech Solutions Inc.（以下简称“客户”）
        
        Party B: Global Services Ltd. (The "Provider")
        乙方：全球服务有限公司（以下简称“提供商”）
        """
        with open(target_filename, "w", encoding="utf-8") as f:
            f.write(dummy_content)
        print(f"✅ Demo file created: {target_filename}")

    # 4. Read and Preview the content (Safety Check)
    # This ensures the file is readable and shows you the first few lines
    try:
        with open(target_filename, 'r', encoding='utf-8') as f:
            preview = f.read(500) # Read only first 500 characters
        
        print("\n--- 📄 File Content Preview (First 500 chars) ---")
        print(preview)
        print("...\n(End of preview)")
        
    except UnicodeDecodeError:
        print("❌ Error: The file encoding is not UTF-8. Please save your txt file as UTF-8.")
    
    return target_filename

# Execute the function and store the filename in a variable
# We will use 'current_processing_file' in Cell 3
current_processing_file = find_and_preview_file()



# %% [Cell 3]

# Cell 3
import pandas as pd
import re
import datetime  # 新增：用于获取当前时间
def is_contains_chinese(string):
    """
    Check if the string contains any Chinese characters.
    """
    for char in string:
        if '\u4e00' <= char <= '\u9fff':
            return True
    return False

def process_text_smartly(input_file, output_excel):
    # 1. Read lines and remove completely empty lines immediately
    with open(input_file, 'r', encoding='utf-8') as f:
        # strict filtering: remove whitespace and check if anything remains
        raw_lines = [line.strip() for line in f.readlines() if line.strip()]

    source_lines = [] # English
    target_lines = [] # Chinese

    # 2. Classify each line based on content, not position
    for line in raw_lines:
        if is_contains_chinese(line):
            target_lines.append(line)
        else:
            source_lines.append(line)

    # 3. Validation
    # We ensure both lists have the same length for the table
    len_source = len(source_lines)
    len_target = len(target_lines)
    
    print(f"Found {len_source} English lines and {len_target} Chinese lines.")
    
    # Handle mismatches by truncating to the shorter length to prevent errors
    min_len = min(len_source, len_target)
    
    df = pd.DataFrame({
        'Source (English)': source_lines[:min_len],
        'Target (Chinese)': target_lines[:min_len]
    })

    df.to_excel(output_excel, index=False)
    print(f"Excel file created: {output_excel}")
    return df

# 1. 获取当前时间并格式化（格式：年月日_时分秒，无特殊字符，兼容全系统）
current_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
# 2. 拼接时间戳到文件名中，保持.xlsx扩展名
output_filename = f"aligned_{current_time}.xlsx"

# Run the function
df_result = process_text_smartly("CGTN-bilingual-news.txt", output_filename)
df_result.head()



# %% [Cell 4]

# Cell 4
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_formatted_word(dataframe, output_filename):
    doc = Document()
    
    # Title
    heading = doc.add_heading('Bilingual Translation Table', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create Table
    table = doc.add_table(rows=len(dataframe) + 1, cols=2)
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Source Text'
    hdr_cells[1].text = 'Target Text'

    # Fill Data
    for i in range(len(dataframe)):
        # Get text from dataframe
        en_text = dataframe.iloc[i]['Source (English)']
        zh_text = dataframe.iloc[i]['Target (Chinese)']
        
        # Get table cells
        row_cells = table.rows[i+1].cells
        
        # Format English Column (Left)
        p_en = row_cells[0].paragraphs[0]
        run_en = p_en.add_run(en_text)
        run_en.font.name = 'Times New Roman'
        run_en.font.size = Pt(11)
        
        # Format Chinese Column (Right)
        p_zh = row_cells[1].paragraphs[0]
        run_zh = p_zh.add_run(zh_text)
        run_zh.font.name = 'SimSun' # Set to Songti
        run_zh.font.size = Pt(11)
        # XML fix for Chinese font rendering
        run_zh._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.save(output_filename)
    print(f"Word document saved as: {output_filename}")
output_docx_filename = f"Layout_{current_time}.docx"
# Run the function
create_formatted_word(df_result, output_docx_filename)

